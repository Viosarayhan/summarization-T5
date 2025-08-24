import streamlit as st
import pandas as pd
from transformers import T5Tokenizer, T5ForConditionalGeneration
import PyPDF2
from docx import Document
import io

st.set_page_config(page_title="Text Summarization App", layout="centered")



MODEL_PATH = "model"

@st.cache_resource
def load_model():
    tokenizer = T5Tokenizer.from_pretrained(MODEL_PATH)
    model = T5ForConditionalGeneration.from_pretrained(MODEL_PATH)
    return tokenizer, model

tokenizer, model = load_model()


def fix_quotes(text):

    count_double = text.count('"')
    count_single = text.count("'")


    if count_double % 2 != 0:
        if text.endswith('"'):
            text = '"' + text
        else:
            text += '"'
    if count_single % 2 != 0:
        if text.endswith("'"):
            text = "'" + text
        else:
            text += "'"

    return text


def summarize(text, max_len=512, min_len=100):
    inputs = tokenizer.encode("summarize: " + text, return_tensors="pt", max_length=1024, truncation=True)
    outputs = model.generate(
        inputs,
        max_length=max_len,
        min_length=min_len,
        length_penalty=2.0,
        num_beams=4,
        early_stopping=True,
        repetition_penalty=2.5,
        no_repeat_ngram_size=3
    )

    decoded = tokenizer.decode(outputs[0], skip_special_tokens=True)
    cleaned = decoded.replace("<pad>", "").strip()

    # ❌ Hapus semua jenis tanda kutip
    cleaned = cleaned.replace("“", "").replace("”", "")
    cleaned = cleaned.replace("‘", "").replace("’", "")
    cleaned = cleaned.replace('"', "").replace("'", "")

    return cleaned

# Fungsi baca PDF
def read_pdf(file):
    pdf_reader = PyPDF2.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

# Fungsi baca DOCX
def read_docx(file):
    doc = Document(io.BytesIO(file.read()))
    full_text = []

    # Tambahkan semua paragraf biasa
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)

    # Tambahkan semua isi tabel
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text for cell in row.cells)
            full_text.append(row_text)

    return "\n".join(full_text)

# UI
st.title("📄 Abstractive Text Summarization")
st.markdown("Model: `T5` hasil fine-tuning pada dataset berita Bahasa Indonesia.")

option = st.radio("📥 Pilih Metode Input:", ["Ketik / Salin Teks", "Upload File (.txt / .csv / .pdf / .docx)"])

text_input = ""

if option == "Ketik / Salin Teks":
    text_input = st.text_area("Masukkan teks artikel di bawah ini:", height=300)

elif option == "Upload File (.txt / .csv / .pdf / .docx)":
    uploaded_file = st.file_uploader("Unggah file artikel", type=["txt", "csv", "pdf", "docx"])
    if uploaded_file is not None:
        if uploaded_file.name.endswith(".txt"):
            text_input = uploaded_file.read().decode("utf-8")
            st.text_area("Isi file .txt:", text_input, height=300)
        elif uploaded_file.name.endswith(".csv"):
            df = pd.read_csv(uploaded_file)
            st.dataframe(df.head())
            col_names = df.columns.tolist()
            selected_col = st.selectbox("Pilih kolom teks untuk diringkas:", col_names)
            text_input = df[selected_col].astype(str).str.cat(sep=" ")[:3000]
        elif uploaded_file.name.endswith(".pdf"):
            try:
                text_input = read_pdf(uploaded_file)
                st.text_area("Teks yang diekstrak dari PDF:", text_input, height=300)
            except Exception as e:
                st.error(f"Gagal membaca file PDF: {e}")
        elif uploaded_file.name.endswith(".docx"):
            try:
                text_input = read_docx(uploaded_file)
                st.text_area("Teks yang diekstrak dari DOCX:", text_input, height=300)
            except Exception as e:
                st.error(f"Gagal membaca file DOCX: {e}")

# Slider panjang ringkasan
col1, col2 = st.columns(2)
with col1:
    max_len = st.slider("🔠 Panjang Maksimum Ringkasan (token)", 64, 512, 256, step=16)
with col2:
    min_len = st.slider("🔡 Panjang Minimum Ringkasan (token)", 32, 256, 100, step=16)

if st.button("🔍 Ringkas Artikel"):
    if not text_input.strip():
        st.warning("Teks tidak boleh kosong.")
    else:
        with st.spinner("🔄 Sedang merangkum..."):
            summary = summarize(text_input, max_len=max_len, min_len=min_len)
            num_tokens = len(tokenizer.encode(summary))
            num_words = len(summary.split())
        st.success("✅ Hasil Ringkasan:")
        st.write(summary)
        st.caption(f"🔢 Panjang ringkasan: {num_tokens} token (~{num_words} kata)")

# Footer
st.markdown("---")
st.markdown("📌 Dibuat menggunakan model T5 hasil fine-tuning pada dataset Liputan6 (Bahasa Indonesia).")
